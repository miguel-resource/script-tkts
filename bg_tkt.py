from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
import time

def create_document(name, description, num_bugs):
    document = Document()
    document.add_heading(f'Bug [{name}]', 0)

    if document is None: print("nop :c")
    add_description(description, document)

    for i in range(0, num_bugs):
        message = f"\nINCIDENCIA #{i + 1} \n"
        print(message)
        before_debugg = input("*** Ingresa el estado del bug antes de debuggear: ")
        after_debugg = input("*** Ingresa el estado de un bug después de debuggear: "),
        imgs_num_before = int(input("*** Ingresa el número de imagenes que tienes como evidencia antes del debugg (número): "))
        imgs_num_after = int (input("*** Ingresa el número de imagenes que tienes como evidencia después de debugg (número): "))
        create_bugs(
            i,
            name,
            document,
            before_debugg,
            after_debugg,
            imgs_num_before,
            imgs_num_after
        )
    document.save(f"files/Bug-fixed[{name}].docx")
    print("yep")


def create_bugs(count, name, document, before_debugg, after_debugg, imgs_num_before, imgs_num_after):
    ## protocol to name images before: './imgs/before/img-bug[{number_tkt}]-{number_incidencia}-{img}.png'

    document.add_heading(f"Incidencia #{count + 1}",1)
    document.add_heading(f"Antes de debuggear:", 2)
    p_before = document.add_paragraph()
    p_before.add_run(before_debugg).font.highlight_color = WD_COLOR_INDEX.RED

    for i in range(0, imgs_num_before):
        new_picture = f"imgs/before/img-bug[{name}]-{count+1}-{i+1}.png"
        document.add_picture(new_picture, width=Inches(5.8))
        print(f"====> añadida {new_picture}")
        time.sleep(0.9)
    document.add_page_break()

    document.add_heading(f"Después de debuggear:", 2)
    p_after = document.add_paragraph()
    p_after.add_run(after_debugg).font.highlight_color = WD_COLOR_INDEX.GREEN

    for e in range(0, imgs_num_after):
        new_picture = f"imgs/fixed/img-bug[{name}]-{count+1}-{e+1}.png"
        document.add_picture(new_picture, width=Inches(5.8))
        print(f"====> añadida {new_picture}")
        time.sleep(0.9)
    document.add_page_break()


def add_description(description, document):
    p = document.add_paragraph()
    p.add_run('Descripción: ').bold = True
    p.add_run(description).font.highlight_color = WD_COLOR_INDEX.YELLOW
    p.add_run('.')


def run():
    message = """
        HOLA, BIENVENIDO A TU GENERADOR DE TKT'S
    """
    print(message)
    num_tkt = input("*** Número de tkt: ")
    description = input("*** Agrega un descripción: ")                                                                                                                                                                                                                                      
    try:
        num_bugs = int(input("*** Agrega el número de incidencias en tu tkt: "))                                                    
    except ValueError:
        print(ValueError)
    create_document(num_tkt, description, num_bugs)



if __name__ == '__main__':
    run()
